from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "paper" / "ProvTimeRAG_WSDM2027_中文论文初稿_v1.docx"

NAVY = "16324F"
BLUE = "2F6B9A"
TEAL = "2A8C82"
LIGHT = "EAF2F7"
PALE = "F5F8FA"
GOLD = "C79A3B"
WHITE = "FFFFFF"
GRAY = "5C6770"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, keep=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if keep and node is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_east_asia(run, font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(3.5)
    pf.first_line_indent = Pt(18)

    for style_name, size, color, before, after in (
        ("Heading 1", 15, NAVY, 11, 5),
        ("Heading 2", 11.5, BLUE, 8, 3),
        ("Heading 3", 10, TEAL, 5, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("PROVTIMERAG  ·  WSDM 2027 中文研究初稿")
    set_east_asia(r)
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("匿名稿  |  2026-08-10  |  ")
    set_east_asia(r)
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_field(fp, "PAGE")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("ProvTimeRAG")
    set_east_asia(r)
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("面向时间敏感检索增强生成的发布者身份与出处状态路由")
    set_east_asia(r)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Anonymous Authors  ·  WSDM 2027 Main Track Chinese Draft")
    set_east_asia(r)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(GRAY)

    callout = doc.add_table(rows=1, cols=3)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    widths = [2.1, 2.1, 2.1]
    labels = [
        ("核心任务", "Publisher-visible provenance routing"),
        ("主结果", "FinFact Top-1 93.83% ± 0.59%"),
        ("下游价值", "Citation recall 0.892 → 0.940"),
    ]
    for i, cell in enumerate(callout.rows[0].cells):
        cell.width = Inches(widths[i])
        set_cell_shading(cell, LIGHT if i != 1 else "DDEFEA")
        set_cell_margins(cell, 100, 110, 100, 110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(labels[i][0] + "\n")
        set_east_asia(r)
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        r = p.add_run(labels[i][1])
        set_east_asia(r)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)


def add_abstract(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("摘要")
    set_east_asia(r)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    text = (
        "检索增强生成通常把证据选择建模为语义相关性排序，但真实 Web 证据还具有发布者、访问入口、文档版本、"
        "有效时间与可支持性等状态。同一段文字在被镜像、归档、转引或跨版本复用后，可能与问题高度相关，却不再来自"
        "目标发布者或目标时间状态。本文提出 ProvTimeRAG，将时间敏感 RAG 的证据控制形式化为出处状态路由。我们首先"
        "建立可观测发布者身份契约，将访问 URL 与内容发布者分离，并通过平衡 donor 的 Source-Swap 构造困难负例；随后"
        "训练一个多任务 cross-encoder，同时学习发布者身份、时间/版本选择和证据不足拒答；在此基础上，使用冻结的保守"
        "结构化解码器进行 bundle 级一致性校正，并在固定生成器与公平候选顺序下评估引用质量。三随机种子实验表明，"
        "C2 在泄漏受控的 clean blind 上取得 88.81%±0.60% Top-1，在 FinFact 跨数据集 11,164 个原子组上取得"
        "93.83%±0.59%，显著优于 Qwen3-Reranker-0.6B 等冻结强基线。移除 URL/domain 观察会在三个评测域下降"
        "15.16–20.17 个百分点。端到端实验中，C3 将引用精确率由 0.500 提升到 0.547、召回率由 0.892 提升到 0.940，"
        "而答案 F1 的差异不显著。结果说明：对时间敏感 Web RAG，显式出处状态建模能够弥补通用相关性排序器的系统性盲点，"
        "其最稳定的下游收益是更可验证的引用，而非无条件提高答案得分。"
    )
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.right_indent = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    set_cell = doc.add_table(rows=1, cols=1).rows[0].cells[0]
    set_cell_shading(set_cell, PALE)
    set_cell_margins(set_cell, 75, 100, 75, 100)
    pp = set_cell.paragraphs[0]
    pp.paragraph_format.first_line_indent = Pt(0)
    rr = pp.add_run("关键词：时间敏感 RAG；出处路由；发布者身份；证据版本；结构化解码；引用可验证性")
    set_east_asia(rr)
    rr.font.size = Pt(8.5)
    rr.font.bold = True
    rr.font.color.rgb = RGBColor.from_string(BLUE)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    return p


def body(doc, text, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_east_asia(r)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_east_asia(r)
    else:
        r = p.add_run(text)
        set_east_asia(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-9)
    r = p.add_run(text)
    set_east_asia(r)
    return p


def add_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            cell.width = Inches(widths[i])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(str(header))
        set_east_asia(r)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(font_size)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if widths:
                cell.width = Inches(widths[i])
            set_cell_margins(cell)
            if ridx % 2:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_east_asia(r)
            r.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_pipeline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("图 1  ProvTimeRAG 的冻结式研究流程")
    set_east_asia(r)
    r.font.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    items = [
        ("C1", "元数据恢复\n人工仲裁"),
        ("→", ""),
        ("C2", "多任务出处\n状态路由"),
        ("→", ""),
        ("C3", "保守全局\n一致性解码"),
        ("→", ""),
        ("C4", "固定生成器\n引用评测"),
    ]
    widths = [1.15, 0.32, 1.35, 0.32, 1.35, 0.32, 1.25]
    for i, (title, sub) in enumerate(items):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_margins(cell, 110, 50, 110, 50)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title == "→":
            r = p.add_run(title)
            set_east_asia(r)
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(GOLD)
        else:
            set_cell_shading(cell, LIGHT if title != "C2" else "DDEFEA")
            r = p.add_run(title + "\n")
            set_east_asia(r)
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(TEAL)
            r = p.add_run(sub)
            set_east_asia(r)
            r.font.size = Pt(8)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(NAVY)


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_abstract(doc)

    heading(doc, "1  引言")
    body(doc, "检索增强生成（RAG）通过外部语料缓解参数知识陈旧和不可追溯问题[1]。然而，标准检索与重排序通常只回答“这段文本是否与问题相关”，没有回答“这段文本现在由谁发布、属于哪个文档版本、在目标时间是否有效、是否真的能够支持该原子结论”。在开放 Web 中，镜像、网页归档、转载、聚合页、社交媒体转引和 PDF 版本共存，使语义相似性与正确出处逐渐分离。对于事实核验、新闻检索和时效性问答，这种分离会把看似合理的错误证据送入生成器。")
    body(doc, "已有 RAG 改进工作主要关注是否检索、检索质量评价、检索后的自我反思或事实修订[6–9]；出处基准则强调证据与回答之间的可追溯性[2,5]。这些方向十分重要，但仍缺少一个面向真实 Web 发布者身份、文档状态和时间状态的统一路由契约。尤其当访问地址是 web.archive.org，而内容发布者是原站点时，直接把 URL host 当作 source 会制造标签错误；当同一文本被换上错误发布者 URL 时，通用 cross-encoder 也可能继续依赖语义捷径。")
    body(doc, "本文提出 ProvTimeRAG。核心观点是：相关性检索之后还需要一个可审计的 provenance-state router。该路由器读取 atomic claim、候选文本、publisher-visible URL/domain、文档元数据和时间线索，联合预测发布者匹配、时间/版本匹配和证据不足。为避免只在内部数据上成立，我们构造 leakage-controlled clean blind，并将冻结模型一次性迁移到 AVerImaTeC 与 FinFact 风格的外部发布者 Source-Swap。")
    body(doc, "本文贡献如下：")
    bullet(doc, "定义 publisher-visible provenance-state routing，并给出访问 URL、实际发布者、文档版本、时间状态和拒答的可复现数据契约。")
    bullet(doc, "提出多任务 C2 路由器，在 clean blind 达到 88.81%±0.60%，在 FinFact 达到 93.83%±0.59%，显著优于四个冻结强 reranker。")
    bullet(doc, "通过 URL/domain、source-only、多任务与外部数据消融识别真正的因果性工程因素，并报告不支持强结论的 C3 与 C1 结果。")
    bullet(doc, "在固定生成器和公平候选顺序下证明，出处路由的稳定收益是引用精确率与召回率，而非无条件提高答案 F1。")

    add_pipeline(doc)

    heading(doc, "2  相关工作")
    heading(doc, "2.1  RAG 与出处可追溯性", 2)
    body(doc, "RAG 将参数模型与非参数语料结合，并把 provenance 视为知识密集任务的重要能力[1]。KILT 进一步统一了多类知识密集任务的知识快照和 provenance 评价[2]。RARR 通过检索与修订为模型输出补充出处[6]；ALCE 系统评估长文本回答的正确性、流畅性和引用质量[7]。与这些工作不同，ProvTimeRAG 位于生成之前，解决候选证据“文本相关但发布者/版本不匹配”的路由问题。")
    heading(doc, "2.2  自适应与纠错式检索", 2)
    body(doc, "Self-RAG 使用反思 token 自适应决定检索与批判[8]，CRAG 使用检索评价器触发纠错动作[9]。这些方法说明检索结果不应被无条件信任，但其评价多聚焦 query–passage relevance 或生成事实性。ProvTimeRAG 把可观测发布者身份和时间/版本状态设为显式监督维度，从而补充相关性评价器的盲区。")
    heading(doc, "2.3  时间知识与事实核验", 2)
    body(doc, "StreamingQA 和 FreshQA 表明模型需要适应持续变化的新闻知识和快速更新事实[10,11]。FEVER 与 AVeriTeC 分别推动大规模证据核验和真实 Web 证据问答[3,4]。AVeriTeC 特别强调不得使用晚于 claim 的证据，这与本文的时间泄漏控制一致。本文进一步强调：即使证据文本内容正确，如果发布者或版本状态不匹配，系统仍缺少可审计的出处保证。")
    heading(doc, "2.4  与通用 reranker 的关系", 2)
    body(doc, "通用 cross-encoder 与指令 reranker 是强相关性基线。本文不将 backbone 本身作为创新，而是固定输入可见性、负例契约、任务组合、平局处理和外部冻结协议，并与 MiniLM、BGE、mxbai 与 Qwen3-Reranker-0.6B 在完全相同的 publisher-visible 输入上比较。")

    heading(doc, "3  问题定义")
    body(doc, "给定问题 q 和原子 claim 集 C={c₁,…,cₙ}，每个 claim cᵢ 对应候选证据集合 Eᵢ。候选 e 包含文本 x、可见访问 URL u、解析后的发布者 p、文档标识 d、时间/版本元数据 τ 和来源角色 r。目标是对每个 cᵢ 选择 e*，或在没有充分证据时输出 abstain。多个 claim 构成 bundle，bundle-level 解码还需控制预测来源集合与跨 claim 一致性。")
    body(doc, "与标准相关性排序不同，正例要求候选同时满足：文本支持 claim；publisher 与目标来源一致；文档/版本状态与问题一致；时间证据不晚于允许边界。Source-Swap 困难负例保持证据文本不变，只替换为另一个合法发布者的 URL 与 publisher identity，因此迫使模型利用出处而不是语义捷径。")
    add_table(doc, ["符号", "含义", "主要风险"], [
        ("q, c", "问题与原子结论", "问题拆分或 claim 泄漏"),
        ("x", "候选证据文本", "语义捷径"),
        ("u, p", "访问 URL 与解析发布者", "归档 host 冒充发布者"),
        ("d, τ", "文档/版本与时间状态", "版本漂移、未来证据"),
        ("A", "拒答/证据不足", "跨域阈值失配"),
        ("B", "多 claim bundle", "来源合并或覆盖不足"),
    ], widths=[0.8, 3.0, 2.4])

    heading(doc, "4  方法")
    heading(doc, "4.1  C1：可观测元数据恢复与仲裁", 2)
    body(doc, "C1 从当前页、URL、标题、站点名、署名、文档发行者和显式引文中恢复 publisher/actor 可观测性及其关系。流程保留 capture timestamp、metadata source URL 和 capture method，并允许 abstain。50 条独立样本和 100 条扩展开发样本由选择题式表格人工审核与二次仲裁。C1 的定位是建立数据契约和审计轨迹，而不是把提示式标注包装成高精度主模型。")
    heading(doc, "4.2  C2：多任务出处状态路由器", 2)
    body(doc, "C2 以 BGE-reranker-v2-m3 为 cross-encoder backbone。每个 candidate 的输入拼接问题、原子 claim、publisher-visible URL/domain、source role、文档元数据和证据文本。模型共享编码器并学习 route、source、temporal、version 与 abstention 目标。训练以 source、temporal/version 和 insufficient 三类组按任务配额采样；最终实验使用 seeds 42/43/44、固定最大长度与相同优化日程。")
    body(doc, "多任务价值不表述为“publisher 一定更高”。计算匹配的 source-only 模型在 publisher-only clean blind 上可略高，但其 temporal/version 能力降至约 0.59–0.76，insufficiency F1 为 0。C2 的目标是以单一 checkpoint 保留接近 source-only 的 publisher 性能，同时完整覆盖时间、版本和拒答。")
    heading(doc, "4.3  C3：冻结的保守结构化解码", 2)
    body(doc, "C3 在候选独立得分之上构造 bundle-level assignment，特征包括 route score 和、来源数量、多来源指示、同来源/同文档/同角色对数以及来源多样性增益。策略只在 train calibration 上拟合，并通过 max-regret 与 source-count-delta 约束选择保守校正。外部测试前冻结 policy ID，测试后不重新选择。")
    heading(doc, "4.4  C4：固定生成与引用契约", 2)
    body(doc, "Raw、C2 与 C3 使用相同 claim group、相同生成器、相同输出 schema 和 label-independent stable-hash 候选顺序。评价 answer token F1、citation precision/recall/hit、abstention、latency 和 token 数。一个供应商拒绝的请求在三种方法中一致排除，并保留原始失败记录。")

    heading(doc, "5  实验设置")
    heading(doc, "5.1  数据与泄漏控制", 2)
    body(doc, "内部 publisher 数据经 v3 clean 合同过滤，train 含 6,655 个 source 组，development 含 460 组，clean blind 含 828 组。train 与 clean blind 的 query overlap 为 0，document overlap 仅 5（清洗后进一步控制），publisher overlap 单独拆分 seen/unseen 报告。FinFact 外部集在冻结前构建，含 2,365 bundles、11,164 groups、22,328 candidates 和 2,850 publishers；正负发布者直方图完全相等，每个 donor 只使用一次。")
    body(doc, "AVerImaTeC 的初版 donor 选择把 163/164 个负例集中到同一域名，因此其 v1 模型结果全部判为无效；v2 用确定性完美匹配平衡 donor。类似地，legacy Source-Swap 的 donor URL host 缺失，修复后才允许进行 URL-aware claim。所有修复都不改变标签、证据文本、group 成员或候选的 label-independent 顺序。")
    add_table(doc, ["数据划分", "Bundles/Groups", "用途", "冻结规则"], [
        ("Publisher train", "6,655 groups", "C2 训练", "query/document leakage controlled"),
        ("Development", "460 groups", "模型开发", "三 seed，同一日程"),
        ("Clean blind", "828 groups", "内部盲测", "checkpoint 冻结后一次评估"),
        ("FinFact", "2,365 / 11,164", "跨数据集外部测试", "balanced donors；测试前冻结"),
        ("C4 fair-order", "678 records", "下游生成", "同候选、同生成器、同顺序"),
    ], widths=[1.3, 1.1, 1.4, 2.8])
    heading(doc, "5.2  基线与实现", 2)
    body(doc, "基线包括 random/BM25、MS MARCO MiniLM-L6-v2、BGE-reranker-v2-m3、mxbai-rerank-large-v1 和 Qwen3-Reranker-0.6B。所有强 reranker 不在测试集拟合、不做模型选择，使用相同 instruction、publisher-visible 输入和 label-independent tie rule。C2 以 FP32 评估，训练可使用 BF16。")
    heading(doc, "5.3  指标与统计", 2)
    body(doc, "原子组使用 Top-1 与 MRR；bundle 使用 exact match 和 publisher source-set exact/Jaccard；拒答报告 accuracy、precision、recall 和 F1。三 seed 报告均值与样本标准差，不把重复 seed 当独立样本合并。同一 seed 内使用 exact McNemar 检验与 paired bootstrap 95% CI[14,15]。")

    heading(doc, "6  结果")
    heading(doc, "6.1  主结果：跨域发布者路由", 2)
    add_table(doc, ["方法", "Clean blind Top-1", "FinFact Top-1", "FinFact bundle exact"], [
        ("MiniLM-L6-v2", "0.5399", "0.5206", "0.1167"),
        ("BGE-reranker-v2-m3", "0.5990", "0.5962", "0.1852"),
        ("mxbai-rerank-large-v1", "0.6087", "0.6745", "0.2837"),
        ("Qwen3-Reranker-0.6B", "0.7343", "0.8793", "0.5979"),
        ("ProvTimeRAG C2", "0.8881 ± 0.0060", "0.9383 ± 0.0059", "0.7686 ± 0.0210"),
        ("C2 + frozen C3", "—", "0.9383（group）", "0.7721 ± 0.0212"),
    ], widths=[2.15, 1.45, 1.35, 1.55])
    body(doc, "C2 在 clean blind 和 FinFact 上分别超过最强冻结 Qwen3 基线约 15.38 和 5.89 个百分点。FinFact 的 C2-versus-Qwen paired p-value 在三个 seed 上分别约为 9.21×10⁻⁷⁸、1.11×10⁻⁶⁴ 和 9.87×10⁻⁴⁹，bootstrap 区间均为正。unseen publisher 的 clean blind 均值为 0.8929±0.0112，说明改进不依赖只记忆训练域名。")
    heading(doc, "6.2  URL/domain 可见性消融", 2)
    add_table(doc, ["划分", "完整 C2", "移除 URL/domain", "下降"], [
        ("Development", "0.8638 ± 0.0082", "0.6862 ± 0.0292", "17.75 pp"),
        ("Clean blind", "0.8881 ± 0.0060", "0.6864 ± 0.0135", "20.17 pp"),
        ("FinFact", "0.9383 ± 0.0059", "0.7866 ± 0.0391", "15.16 pp"),
    ], widths=[1.7, 1.65, 1.75, 1.2])
    body(doc, "该消融是论文最直接的机制证据：保持训练框架和任务不变，仅移除 publisher-visible URL/domain 就在三个域稳定下降。FinFact 上 no-URL 模型还把约 97% 的可回答组误判为 abstain，暴露出跨域 calibration collapse。我们不在 FinFact 重校准阈值，因为那会污染冻结外部测试。")
    heading(doc, "6.3  多任务与 source-only 对照", 2)
    body(doc, "compute-matched source-only 三 seed 在 development 上约 0.8703，在 clean blind 上约 0.8961，在 FinFact 上约 0.9376；C2 对应为 0.8638、0.8881、0.9383。差异很小且方向随域变化，不能声称多任务显著提高 publisher。关键差别是 source-only 在 temporal/version 上明显退化、insufficiency F1=0，而 C2 的 temporal/version 接近 1.0、insufficiency F1=1.0。数据匹配而非计算匹配的 source-only 结果将在最终英文稿中作为补充稳定性检查；在其同步前不填写未知数值。")
    heading(doc, "6.4  C3：小幅、保守但不显著", 2)
    body(doc, "冻结 C3 将 FinFact bundle exact 从 0.7686 提升到 0.7721，三 seed 平均 +0.35 pp；McNemar p=0.233/0.280/0.280，paired bootstrap CI 均跨 0。publisher source-set Jaccard 有轻微改善，但 binary multi-source error 从 0.00747 上升到 0.00888。由此，C3 只作为可解释的保守解码器和消融，不作为主显著性结论。")
    heading(doc, "6.5  C4：下游引用质量", 2)
    add_table(doc, ["输入策略", "Answer F1", "Citation P", "Citation R", "Citation hit"], [
        ("Raw", "0.6192", "0.4998", "0.8916", "0.9071"),
        ("C2", "0.6162", "0.5302", "0.9028", "0.9189"),
        ("C3", "0.6143", "0.5470", "0.9400", "0.9543"),
    ], widths=[1.2, 1.15, 1.2, 1.2, 1.2])
    body(doc, "C3 相对 Raw 的 citation precision、recall 与 hit 提升的 paired bootstrap 95% CI 均完全大于 0；answer F1 差异区间跨 0。更细的错误分析显示，单证据问题上的答案 F1 基本持平，而多证据问题仍是主要困难。这支持“出处路由改善可验证引用”的结论，也说明未来不应通过更激进裁剪牺牲多证据覆盖。")

    heading(doc, "7  分析与讨论")
    heading(doc, "7.1  为什么通用 reranker 不够", 2)
    body(doc, "Source-Swap 保持文本不变，只改变发布者身份；因此只依赖 claim–text 语义的模型接近二选一。Qwen3-Reranker 能利用 instruction 和元数据，FinFact 达到 0.8793，但仍落后于显式 publisher 监督的 C2。该差距不是 backbone 大小本身，而是监督目标、输入可见性与困难负例共同作用。")
    heading(doc, "7.2  C1 的真实作用", 2)
    body(doc, "C1 独立 50 条人工金标的最佳 row exact 为 0.48，字段准确率分布不均。该数值不足以支撑“自动元数据标注已解决”，但它揭示 publisher/actor relation 和 attribution link 的歧义，并提供了固定选项、仲裁和捕获来源的审计机制。论文应把 C1 放在数据治理与可行性分析中，而非主结果表。")
    heading(doc, "7.3  失败模式", 2)
    bullet(doc, "归档与镜像：访问 host 与真实发布者不一致，必须显式解包或拒绝。")
    bullet(doc, "多证据问题：激进路由提高引用精确率时可能降低证据覆盖和答案 F1。")
    bullet(doc, "跨域拒答：没有 publisher 可见性时，abstention logit 可能发生大幅分布漂移。")
    bullet(doc, "结构化策略：即使平均 bundle exact 为正，也可能增加来源数量错误，必须多指标约束。")

    heading(doc, "8  局限性")
    body(doc, "第一，Publisher identity 是 operational identity，不等同于法律责任主体或事实可信度；模型判断“来源匹配”不代表来源真实。第二，外部集来自事实核验数据的文本证据，图像本身不进入模型，因此不覆盖多模态出处。第三，C1 人工金标规模有限，且 relation 标签存在主观边界。第四，C4 使用一个冻结 API 生成器，结论主要关于引用行为，未证明对所有 LLM 泛化。第五，FinFact 构造虽严格平衡 donor，但仍可能保留主题或站点风格线索；未来应加入跨语言、跨时间和对抗式 publisher spoofing。")

    heading(doc, "9  伦理考虑与可复现性")
    body(doc, "本研究处理公开 Web 证据和事实核验数据，不发布 API 密钥、私人注释、原始抓取缓存或受许可约束的模型权重。公开仓库只提供代码、小型结果摘要、哈希和构建说明。URL 与网页文本可能包含个人信息或已删除内容，数据发布应遵守原许可、robots/服务条款和必要的最小化原则。系统可能因错误 publisher 解析而误归因，不应直接用于法律、新闻发布或高风险自动裁决。所有外部结果在测试前冻结构造与 policy；失败请求按预先声明规则一致处理。")

    heading(doc, "10  结论")
    body(doc, "ProvTimeRAG 将时间敏感 RAG 的证据选择从纯相关性排序推进到 publisher-visible provenance-state routing。严格的数据修复、泄漏控制、平衡 Source-Swap、三 seed 外部评测和强 reranker 对比表明，显式发布者身份是稳定且可泛化的信号。多任务 C2 以一个 checkpoint 联合覆盖 publisher、时间/版本和拒答；C3 的外部增益小且不显著，但 C4 显示其能够显著改善引用精确率与召回率。整体证据支持一项克制但重要的结论：在开放 Web RAG 中，可靠出处需要被建模，而不能假设语义相关性会自动携带 provenance。")

    heading(doc, "参考文献")
    refs = [
        "[1] Lewis, P. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.",
        "[2] Petroni, F. et al. KILT: a Benchmark for Knowledge Intensive Language Tasks. NAACL, 2021.",
        "[3] Thorne, J. et al. FEVER: a Large-scale Dataset for Fact Extraction and Verification. NAACL, 2018.",
        "[4] Schlichtkrull, M., Guo, Z., and Vlachos, A. AVeriTeC: A Dataset for Real-world Claim Verification with Evidence from the Web. NeurIPS Datasets and Benchmarks, 2023.",
        "[5] Gao, T., Yen, H., Yu, J., and Chen, D. Enabling Large Language Models to Generate Text with Citations. EMNLP, 2023.",
        "[6] Gao, L. et al. RARR: Researching and Revising What Language Models Say, Using Language Models. ACL, 2023.",
        "[7] Huang, C. et al. Training Language Models to Generate Text with Citations via Fine-grained Rewards. ACL, 2024.",
        "[8] Asai, A. et al. Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection. ICLR, 2024.",
        "[9] Yan, S.-Q., Gu, J.-C., Zhu, Y., and Ling, Z.-H. Corrective Retrieval Augmented Generation. arXiv:2401.15884, 2024.",
        "[10] Liska, A. et al. StreamingQA: A Benchmark for Adaptation to New Knowledge over Time in Question Answering Models. ICML, 2022.",
        "[11] Vu, T. et al. FreshLLMs: Refreshing Large Language Models with Search Engine Augmentation. arXiv:2310.03214, 2023.",
        "[12] Chen, J. et al. BGE M3-Embedding: Multi-Lingual, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation. arXiv:2402.03216, 2024.",
        "[13] Zhang, Y. et al. Qwen3 Embedding: Advancing Text Embedding and Reranking Through Foundation Models. arXiv:2506.05176, 2025.",
        "[14] McNemar, Q. Note on the Sampling Error of the Difference Between Correlated Proportions or Percentages. Psychometrika, 1947.",
        "[15] Efron, B. Bootstrap Methods: Another Look at the Jackknife. The Annals of Statistics, 1979.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.first_line_indent = Pt(-16)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ref)
        set_east_asia(r)
        r.font.size = Pt(8)

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "附录 A  投稿迁移与剩余核验清单")
    body(doc, "本 Word 是中文研究初稿，不是最终投稿格式。正式 WSDM 2027 主会稿须翻译为英文，使用 ACM sigconf,anonymous,review 模板，正文（含图表和附录）不超过 9 页，参考文献与伦理声明不受同一页数限制。")
    add_table(doc, ["事项", "当前状态", "写作处理"], [
        ("C2 三 seed + 外部强基线", "已完成", "主表与显著性"),
        ("URL/domain 消融", "已完成", "核心机制表"),
        ("compute-matched source-only", "已完成", "克制表述多任务价值"),
        ("data-matched source-only", "服务器结果待同步", "仅补充稳定性检查"),
        ("C3 frozen external", "已完成；不显著", "次要分析，不夸大"),
        ("C4 fair-order 678", "已完成", "引用质量主证据"),
        ("C1 人工仲裁", "已完成", "可行性/数据治理"),
        ("GitHub remote", "尚未配置", "本地提交后绑定并推送"),
    ], widths=[2.0, 1.8, 2.6])
    heading(doc, "附录 B  推荐优先阅读")
    body(doc, "为了快速进入英文写作，建议按以下顺序阅读：AVeriTeC（真实 Web 事实核验和时间泄漏）；KILT（provenance 评价）；ALCE（引用指标与实验写法）；RARR（归因与修订）；Self-RAG/CRAG（检索质量与自适应路由）；StreamingQA/FreshLLMs（时间变化）。阅读时重点模仿问题动机、评测合同和 limitation 写法，而不是机械复制方法结构。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "ProvTimeRAG：面向时间敏感检索增强生成的发布者身份与出处状态路由"
    doc.core_properties.subject = "WSDM 2027 Chinese research draft"
    doc.core_properties.author = "Anonymous Authors"
    doc.core_properties.keywords = "RAG, provenance, publisher identity, temporal routing, citations"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
